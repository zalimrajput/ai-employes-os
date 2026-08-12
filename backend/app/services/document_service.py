"""Document intelligence: extract text, chunk, embed and store for RAG.

Supported types: txt/md/csv/json (PEP text), pdf, docx, and xlsx (via
openpyxl).  Text is chunked with a simple sliding window and stored as
KnowledgeArticle rows with a pgvector embedding; the original file metadata is
kept in the documents table so uploads remain visible in the UI.
"""
import logging
import re
from typing import Optional

from sqlalchemy.orm import Session

from app.ai.embeddings import embed
from app.models.document import Document
from app.models.knowledge_base import KnowledgeArticle
from app.services.ocr_service import is_image_filename, needs_ocr_fallback

logger = logging.getLogger("app.services.document_service")

CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200

_TEXT_EXTENSIONS = {"txt", "md", "markdown", "csv", "json", "log", "html", "htm"}


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\s+", " ", text or "").strip()
    if len(text) <= size:
        return [text] if text else []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def extract_text(filename: str, raw: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        text = _extract_text_from_pdf_bytes(raw)
        if needs_ocr_fallback(text):
            ocr = _ocr_text_from_pdf_bytes(raw)
            if ocr:
                return ocr
        return text
    if is_image_filename(name):
        ocr_text = _ocr_text_from_image_bytes(raw)
        if ocr_text:
            return ocr_text
    if name.endswith(".docx"):
        from docx import Document as DocxDocument
        from io import BytesIO

        doc = DocxDocument(BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    if name.endswith(".xlsx") or name.endswith(".xls"):
        from io import BytesIO

        from openpyxl import load_workbook

        wb = load_workbook(BytesIO(raw), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                lines.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(lines)
    # Plain text family
    try:
        return raw.decode("utf-8", errors="replace")
    except Exception:
        return raw.decode("latin-1", errors="replace")


def _extract_text_from_pdf_bytes(raw: bytes) -> str:
    from io import BytesIO

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(raw))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _ocr_text_from_pdf_bytes(raw: bytes) -> str:
    from app.services.ocr_service import extract_text_from_pdf

    return extract_text_from_pdf(raw)


def _ocr_text_from_image_bytes(raw: bytes) -> str:
    from app.services.ocr_service import extract_text_from_image

    return extract_text_from_image(raw)


def ingest_document(
    db: Session,
    organization_id,
    uploaded_by,
    filename: str,
    raw: bytes,
    title: str | None = None,
    source: str = "upload",
) -> dict:
    """Extract, chunk (optional) and index a document.

    Creates:
      - a ``documents`` row (metadata + extracted text; embedding on the whole
        doc when a key is available)
      - one or more ``knowledge_articles`` chunks (each embedded) so RAG search
        finds relevant fragments.
    Returns summary dict for logging/response.
    """
    text = extract_text(filename, raw)
    doc = Document(
        organization_id=organization_id,
        uploaded_by=uploaded_by,
        filename=filename,
        mime_type=_mime_for(filename),
        size=len(raw),
        extracted_text=text[:400_000],
    )
    db.add(doc)
    db.flush()

    chunks = chunk_text(text)[:50]
    vectors = embed(chunks) if chunks else None
    if chunks and vectors is None:
        logger.warning(
            "embeddings not created for %s — no embedding provider configured",
            filename,
        )

    doc.embedding = vectors[0] if vectors else _embedding_for_chunk(text)
    articles = []
    for i, chunk in enumerate(chunks):
        article = KnowledgeArticle(
            organization_id=organization_id,
            title=f"{title or filename} — part {i + 1}",
            content=chunk,
            source=source,
        )
        article.embedding = _safe_vector(vectors, i)
        db.add(article)
        articles.append(article)
    db.commit()
    for a in articles:
        db.refresh(a)
    db.refresh(doc)
    return {
        "document_id": str(doc.id),
        "articles": [str(a.id) for a in articles],
        "extracted_chars": len(text),
        "chunks": len(chunks),
    }


def _safe_vector(vectors, index):
    if vectors and index < len(vectors):
        return vectors[index]
    return None


def _embedding_for_chunk(chunk: str) -> list[float]:
    embeds = embed([chunk])
    return embeds[0] if embeds else None


def _mime_for(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "csv": "text/csv",
        "md": "text/markdown",
    }.get(ext, "application/octet-stream")


# Backwards-compat alias for the worker module.
chunk_document = chunk_text